"""
Resume Generator Service for creating PDF and DOCX resumes.
"""

import os
import re
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem, KeepTogether
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models.resume import ResumeData
from core.config import settings


class ResumeGeneratorService:
    """Service for generating ATS-friendly PDF and DOCX resumes."""

    @staticmethod
    def _ensure_generated_dir():
        """Ensure the generated directory exists."""
        os.makedirs(settings.GENERATED_DIR, exist_ok=True)

    @staticmethod
    def _build_filename(resume_data: "ResumeData", ext: str) -> str:
        """
        Build a human-readable, filesystem-safe filename.
        Format: <ContactName>_<JobTitle>.<ext>
        Example: John_Doe_Software_Engineer.pdf
        If a file with the same name already exists it is simply overwritten.
        """
        # --- Contact name ---
        raw_name = resume_data.contact_info.get("name", "").strip()
        # --- Most recent job title (first experience entry after tailoring) ---
        raw_title = ""
        if resume_data.experience:
            raw_title = resume_data.experience[0].title.strip()

        def slugify(text: str) -> str:
            """Replace spaces with underscores and strip unsafe characters."""
            text = re.sub(r"[^\w\s\-]", "", text)   # keep word chars, spaces, hyphens
            text = re.sub(r"\s+", "_", text.strip())  # spaces → underscores
            return text[:50]  # cap length per segment

        name_slug = slugify(raw_name) or "Resume"
        title_slug = slugify(raw_title)

        parts = [name_slug]
        if title_slug:
            parts.append(title_slug)

        return "_".join(parts) + f".{ext}"

    @staticmethod
    def _format_date_range(start, end) -> str:
        """
        Format date range for resume.
        Accepts strings like "2020-01", "2020-01-15", or datetime.date/datetime.
        Returns "Mon YYYY – Mon YYYY" or "Mon YYYY – Present".
        """
        def fmt(dt):
            if isinstance(dt, str):
                for fmt_str in ("%Y-%m-%d", "%Y-%m", "%m/%Y", "%b %Y"):
                    try:
                        dt = datetime.strptime(dt, fmt_str)
                        break
                    except ValueError:
                        continue
                else:
                    return dt
            if not isinstance(dt, datetime):
                return str(dt)
            return dt.strftime("%b %Y")

        if not start:
            return ""
        start_str = fmt(start)
        end_str = "Present" if not end else fmt(end)
        if start_str and end_str:
            return f"{start_str} – {end_str}"
        return start_str

    @staticmethod
    def _pdf_hr():
        """Return a full-width horizontal rule for use after PDF section headers."""
        return HRFlowable(
            width="100%",
            thickness=0.75,
            color=colors.black,
            spaceBefore=1,
            spaceAfter=4,
        )

    @staticmethod
    def _docx_bottom_border(paragraph):
        """Add a full-width bottom border line to a DOCX paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")      # border thickness (eighths of a point)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @classmethod
    def generate_pdf(cls, resume_data: ResumeData) -> str:
        """
        Generate an ATS-friendly PDF resume from ResumeData.

        Returns:
            The relative path to the generated PDF file (from backend root).
        """
        cls._ensure_generated_dir()

        filename = cls._build_filename(resume_data, "pdf")
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        MARGIN = 36  # 0.5 inch in points
        doc = SimpleDocTemplate(
            filepath,
            pagesize=LETTER,
            leftMargin=MARGIN,
            rightMargin=MARGIN,
            topMargin=MARGIN,
            bottomMargin=MARGIN,
            title="Optimized Resume",
            author=resume_data.contact_info.get("name", "Candidate"),
        )

        base_font = "Times-Roman"
        base_size = 11
        leading = int(base_size * 1.2)

        styles = getSampleStyleSheet()
        styles["Normal"].fontName = base_font
        styles["Normal"].fontSize = base_size
        styles["Normal"].leading = leading
        styles["Normal"].alignment = TA_LEFT
        styles["Normal"].spaceAfter = 6

        styles.add(
            ParagraphStyle(
                name="Name",
                parent=styles["Normal"],
                fontName=base_font,
                fontSize=20,
                leading=24,
                alignment=TA_CENTER,
                spaceAfter=6,
                textColor=colors.black,
            )
        )
        styles.add(
            ParagraphStyle(
                name="Contact",
                parent=styles["Normal"],
                fontName=base_font,
                fontSize=10,
                leading=12,
                alignment=TA_CENTER,
                spaceAfter=12,
                textColor=colors.black,
            )
        )
        styles.add(
            ParagraphStyle(
                name="SectionHeader",
                parent=styles["Normal"],
                fontName="Times-Bold",
                fontSize=12,
                leading=14,
                spaceBefore=12,
                spaceAfter=2,   # tight gap before the HR rule
                textColor=colors.black,
            )
        )
        styles.add(
            ParagraphStyle(
                name="JobTitle",
                parent=styles["Normal"],
                fontName=base_font,
                fontSize=11,
                leading=13,
                spaceBefore=6,
                spaceAfter=2,
                textColor=colors.black,
            )
        )
        styles.add(
            ParagraphStyle(
                name="Company",
                parent=styles["Normal"],
                fontName=base_font,
                fontSize=10,
                leading=12,
                textColor=colors.darkgray,
            )
        )
        styles.add(
            ParagraphStyle(
                name="DateRange",
                parent=styles["Normal"],
                fontName=base_font,
                fontSize=9,
                leading=11,
                textColor=colors.grey,
            )
        )
        styles.add(
            ParagraphStyle(
                name="BulletStyle",
                parent=styles["Normal"],
                fontName=base_font,
                fontSize=base_size,
                leading=leading,
                leftIndent=12,
                bulletIndent=0,
                bulletFontName=base_font,
                bulletFontSize=base_size,
                bulletOffsetY=2,
                textColor=colors.black,
            )
        )

        flow = []

        # Name
        name = resume_data.contact_info.get("name")
        if name:
            flow.append(Paragraph(name, styles["Name"]))

        # Contact line
        contact_bits = []
        for key in ["email", "phone", "location", "linkedin", "github", "website"]:
            val = resume_data.contact_info.get(key)
            if val:
                contact_bits.append(val)
        if contact_bits:
            flow.append(Paragraph(" | ".join(contact_bits), styles["Contact"]))

        # Professional Summary
        if resume_data.summary:
            flow.append(Paragraph("Professional Summary", styles["SectionHeader"]))
            flow.append(cls._pdf_hr())
            flow.append(Paragraph(resume_data.summary, styles["Normal"]))
            flow.append(Spacer(1, 6))

        # Skills
        if resume_data.skill_categories:
            flow.append(Paragraph("Technical Skills", styles["SectionHeader"]))
            flow.append(cls._pdf_hr())
            for cat_name, cat_skills in resume_data.skill_categories.items():
                flow.append(Paragraph(f"<b>{cat_name}:</b> {', '.join(cat_skills)}", styles["Normal"]))
            flow.append(Spacer(1, 6))
        elif resume_data.skills:
            flow.append(Paragraph("Skills", styles["SectionHeader"]))
            flow.append(cls._pdf_hr())
            flow.append(Paragraph(", ".join(resume_data.skills), styles["Normal"]))
            flow.append(Spacer(1, 6))

        # Experience
        if resume_data.experience:
            flow.append(Paragraph("Experience", styles["SectionHeader"]))
            flow.append(cls._pdf_hr())
            for exp in resume_data.experience:
                job_line = f"<b>{exp.title}</b> – <b>{exp.company}</b>"
                flow.append(Paragraph(job_line, styles["JobTitle"]))
                date_range = cls._format_date_range(exp.start_date, exp.end_date)
                if date_range:
                    flow.append(Paragraph(date_range, styles["DateRange"]))
                if exp.description:
                    bullets = exp.description if isinstance(exp.description, list) else [s.strip() for s in exp.description.split(".") if s.strip()]
                    for bullet in bullets:
                        flow.append(Paragraph(bullet, styles["BulletStyle"], bulletText="•"))
                flow.append(Spacer(1, 6))

        # Projects
        if resume_data.projects:
            flow.append(Paragraph("Projects", styles["SectionHeader"]))
            flow.append(cls._pdf_hr())
            for proj in resume_data.projects:
                flow.append(Paragraph(f"<b>{proj.name}</b>", styles["JobTitle"]))
                if proj.description:
                    bullets = proj.description if isinstance(proj.description, list) else [s.strip() for s in proj.description.split(".") if s.strip()]
                    for bullet in bullets:
                        flow.append(Paragraph(bullet, styles["BulletStyle"], bulletText="•"))
                flow.append(Spacer(1, 6))

        # Education
        if resume_data.education:
            flow.append(Paragraph("Education", styles["SectionHeader"]))
            flow.append(cls._pdf_hr())
            for edu in resume_data.education:
                edu_line = f"<b>{edu.degree}</b> – <b>{edu.institution}</b>"
                flow.append(Paragraph(edu_line, styles["JobTitle"]))
                if getattr(edu, "graduation_year", None):
                    flow.append(Paragraph(str(edu.graduation_year), styles["DateRange"]))
                if getattr(edu, "gpa", None):
                    flow.append(Paragraph(f"GPA: {edu.gpa}", styles["Normal"]))
                flow.append(Spacer(1, 6))

        # Achievements
        if resume_data.achievements:
            flow.append(Paragraph("Achievements", styles["SectionHeader"]))
            flow.append(cls._pdf_hr())
            for ach in resume_data.achievements:
                flow.append(Paragraph(ach, styles["BulletStyle"], bulletText="•"))
            flow.append(Spacer(1, 6))

        doc.build(flow)
        return os.path.join(settings.GENERATED_DIR, filename)

    @classmethod
    def generate_docx(cls, resume_data: ResumeData) -> str:
        """
        Generate an ATS-friendly DOCX resume from ResumeData.

        Returns:
            The relative path to the generated DOCX file (from backend root).
        """
        cls._ensure_generated_dir()

        filename = cls._build_filename(resume_data, "docx")
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        def add_run(paragraph, text, size=None, bold=False, italic=False, color=None):
            run = paragraph.add_run(text)
            if size is not None:
                run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color is not None:
                run.font.color.rgb = color
            run.font.name = 'Times New Roman'
            return run

        def add_section_header(title):
            """Add a bold section heading followed by a full-width bottom border line."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
            add_run(p, title, size=12, bold=True)
            cls._docx_bottom_border(p)
            return p

        # ---- Name (centered, 20pt, bold) ---------------------------
        name = resume_data.contact_info.get("name")
        if name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, name, size=20, bold=True)

        # ---- Contact line (centered, 10pt) -------------------------
        contact_parts = []
        for key in ["email", "phone", "location", "linkedin", "github", "website"]:
            val = resume_data.contact_info.get(key)
            if val:
                contact_parts.append(val)
        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, " | ".join(contact_parts), size=10)

        # ---- Professional Summary ----------------------------------
        if resume_data.summary:
            add_section_header("Professional Summary")
            p = doc.add_paragraph(resume_data.summary)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        # ---- Skills ------------------------------------------------
        if resume_data.skill_categories:
            add_section_header("Technical Skills")
            for cat_name, cat_skills in resume_data.skill_categories.items():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                add_run(p, f"{cat_name}: ", size=11, bold=True)
                add_run(p, ", ".join(cat_skills), size=11)
        elif resume_data.skills:
            add_section_header("Skills")
            skills_text = ", ".join(resume_data.skills)
            p = doc.add_paragraph(skills_text)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        # ---- Experience --------------------------------------------
        if resume_data.experience:
            add_section_header("Experience")
            for exp in resume_data.experience:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                add_run(p, f"{exp.title} – ", size=11, bold=True)
                add_run(p, exp.company, size=11, bold=True)
                date_range = cls._format_date_range(exp.start_date, exp.end_date)
                if date_range:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    add_run(p, date_range, size=9, italic=True, color=RGBColor(85, 85, 85))
                if exp.description:
                    bullets = exp.description if isinstance(exp.description, list) else [s.strip() for s in exp.description.split(".") if s.strip()]
                    for bullet in bullets:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.paragraph_format.left_indent = Inches(0.25)
                        bp.paragraph_format.line_spacing = 1.15
                        add_run(bp, bullet, size=11)
                doc.add_paragraph()

        # ---- Projects ----------------------------------------------
        if resume_data.projects:
            add_section_header("Projects")
            for proj in resume_data.projects:
                p = doc.add_paragraph()
                add_run(p, proj.name, size=11, bold=True)
                if proj.description:
                    bullets = proj.description if isinstance(proj.description, list) else [s.strip() for s in proj.description.split(".") if s.strip()]
                    for bullet in bullets:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.paragraph_format.left_indent = Inches(0.25)
                        bp.paragraph_format.line_spacing = 1.15
                        add_run(bp, bullet, size=11)
                doc.add_paragraph()

        # ---- Education ---------------------------------------------
        if resume_data.education:
            add_section_header("Education")
            for edu in resume_data.education:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                add_run(p, f"{edu.degree} – ", size=11, bold=True)
                add_run(p, edu.institution, size=11, bold=True)
                if getattr(edu, "graduation_year", None):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    add_run(p, str(edu.graduation_year), size=9, italic=True)
                if getattr(edu, "gpa", None):
                    p = doc.add_paragraph()
                    add_run(p, f"GPA: {edu.gpa}", size=11)
                doc.add_paragraph()

        # ---- Achievements ------------------------------------------
        if resume_data.achievements:
            add_section_header("Achievements")
            for ach in resume_data.achievements:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.line_spacing = 1.15
                add_run(p, ach, size=11)
            doc.add_paragraph()

        doc.save(filepath)
        return os.path.join(settings.GENERATED_DIR, filename)
